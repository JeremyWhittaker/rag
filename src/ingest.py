"""PDF + DOCX → vector-store indexer with legal document enhancements
-------------------------------------------------------------------
• Recursively crawls --pdf-dir looking for *.pdf and *.docx
• Extracts text with legal-specific processing:
      PDFs: unstructured.partition_pdf(strategy="fast")  → auto-OCR fallback
      DOCX: UnstructuredWordDocumentLoader
• Legal metadata extraction (citations, parties, dates, case numbers)
• SHA-256 hash dedup: only new or changed files are embedded
• Hierarchical chunking for legal documents (holdings > quotes > general)
• Saves each project in its own Chroma collection (index/<project>/)
"""

from __future__ import annotations   # must stay first!

# ── stdlib ────────────────────────────────────────────────────────────────
import argparse, hashlib, logging, os
from pathlib import Path
from typing import Sequence, Dict, List

logging.getLogger("pdfminer").setLevel(logging.ERROR)   # kill colour-space spam
os.environ["SCARF_NO_ANALYTICS"] = "true"               # disable telemetry

# ── 3rd-party ─────────────────────────────────────────────────────────────
from tqdm import tqdm
from unstructured.partition.pdf import partition_pdf
from langchain_community.document_loaders.word_document import (
    UnstructuredWordDocumentLoader,
)
from langchain.schema import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_chroma import Chroma

# ── local ─────────────────────────────────────────────────────────────────
from .config import INDEX_DIR, EMBED_MODEL
from .logger import get_logger
from .legal_processor import LegalDocumentProcessor
from .legal_metadata import LegalDocumentMetadata, MetadataExtractor, MetadataIndex

LOGGER = get_logger(__name__)

# ── helper funcs ──────────────────────────────────────────────────────────
def sha256(path: Path) -> str:
    h = hashlib.sha256()
    h.update(path.read_bytes())
    return h.hexdigest()


def load_pdf(path: Path, *, strategy: str = "fast") -> Sequence[str]:
    """Return text for a single PDF; 'fast' auto-OCRs image pages."""
    return [
        el.text
        for el in partition_pdf(filename=str(path), strategy=strategy)
        if el.text and not el.text.isspace()
    ]


# ── main ingest ───────────────────────────────────────────────────────────
def ingest(
    *,
    pdf_dir: Path,
    project: str,
    root_index: Path = INDEX_DIR,
    chunk_size: int = 800,
    enable_legal_processing: bool = True,
) -> None:
    proj_dir = root_index / project
    proj_dir.mkdir(parents=True, exist_ok=True)

    db = Chroma(
        persist_directory=str(proj_dir),
        collection_name=project,
        embedding_function=OpenAIEmbeddings(model=EMBED_MODEL),
    )

    existing = {m.get("sha256") for m in db.get()["metadatas"] if "sha256" in m}

    # Initialize legal processors if enabled
    if enable_legal_processing:
        legal_processor = LegalDocumentProcessor()
        metadata_extractor = MetadataExtractor()
        metadata_index = MetadataIndex(proj_dir)
    
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size, chunk_overlap=120
    )
    new_docs: list[Document] = []
    
    files = list(pdf_dir.rglob("*.pdf")) + list(pdf_dir.rglob("*.docx")) + list(pdf_dir.rglob("*.doc"))
    if not files:
        LOGGER.error("No PDF or DOCX files found in %s", pdf_dir)
        return

    for path in tqdm(files, desc="Parsing documents"):
        digest = sha256(path)
        if digest in existing:
            continue

        # Extract text
        if path.suffix.lower() == ".pdf":
            texts = load_pdf(path)
        else:  # .docx / .doc
            try:
                # Try with metadata extraction disabled
                from unstructured.partition.docx import partition_docx
                elements = partition_docx(
                    filename=str(path),
                    metadata_filename=None,
                    include_metadata=False
                )
                texts = [el.text for el in elements if el.text and not el.text.isspace()]
            except Exception as e:
                LOGGER.warning(f"Error loading {path} with unstructured: {e}")
                # Fallback to simple docx extraction
                try:
                    import docx
                    doc = docx.Document(str(path))
                    texts = [para.text for para in doc.paragraphs if para.text and not para.text.isspace()]
                except Exception as e2:
                    LOGGER.error(f"Failed to load {path} with both methods: {e2}")
                    texts = []  # Skip this file
        
        # Combine all text for legal analysis
        full_text = "\n".join(texts)
        
        # Legal document processing
        base_metadata = {"source": path.name, "sha256": digest}
        
        if enable_legal_processing and full_text:
            # Classify document and extract legal information
            doc_info = legal_processor.classify_document(full_text, path.name)
            
            # Extract comprehensive metadata
            legal_metadata = metadata_extractor.extract_from_text(
                full_text, path, digest, doc_info
            )
            
            # Add to metadata index
            metadata_index.add_document(legal_metadata)
            
            # Enhance base metadata with legal information
            base_metadata.update({
                "document_type": legal_metadata.document_type,
                "authority_weight": legal_metadata.authority_weight,
                "is_primary_authority": legal_metadata.is_primary_authority,
                "case_number": legal_metadata.case_number,
                "jurisdiction": legal_metadata.jurisdiction,
            })
            
            # Add citations to metadata for search
            if legal_metadata.statutes_cited:
                base_metadata["statutes_cited"] = ", ".join(legal_metadata.statutes_cited[:5])
            if legal_metadata.cases_cited:
                base_metadata["cases_cited"] = ", ".join(legal_metadata.cases_cited[:5])
            
            # Create hierarchical chunks for important legal content
            priority_chunks = legal_processor.create_hierarchical_chunks(full_text, doc_info)
            
            for chunk_info in priority_chunks:
                chunk_metadata = base_metadata.copy()
                chunk_metadata.update({
                    "chunk_priority": chunk_info["priority"],
                    "chunk_type": chunk_info["type"],
                })
                new_docs.append(
                    Document(
                        page_content=chunk_info["content"],
                        metadata=chunk_metadata,
                    )
                )
        
        # Process regular text chunks
        for text in texts:
            new_docs.append(
                Document(
                    page_content=text,
                    metadata=base_metadata,
                )
            )

    if not new_docs:
        LOGGER.info("No new documents to embed for project '%s'", project)
        return

    LOGGER.info("Splitting into chunks …")
    chunks = splitter.split_documents(new_docs)
    
    # Sort chunks by priority if legal processing is enabled
    if enable_legal_processing:
        chunks.sort(
            key=lambda x: x.metadata.get("chunk_priority", 999)
        )
    
    LOGGER.info("Embedding %d new chunks", len(chunks))

    # Add documents in batches to avoid memory issues
    batch_size = 100
    for i in range(0, len(chunks), batch_size):
        batch = chunks[i:i + batch_size]
        db.add_documents(batch)
    
    # Save metadata index
    if enable_legal_processing:
        try:
            metadata_index.save_index()
            stats = metadata_index.get_statistics()
            LOGGER.info("Legal document statistics: %s", stats)
        except Exception as e:
            LOGGER.error(f"Error saving metadata index: {e}")
    
    LOGGER.info("✅ Finished ingest for %s (stored at %s)", project, proj_dir)


if __name__ == "__main__":
    ap = argparse.ArgumentParser(description="Ingest PDFs & DOCX into a project with legal enhancements")
    ap.add_argument("--pdf-dir", type=Path, required=True)
    ap.add_argument("--project", required=True, help="Project name (court_case, statutes, etc.)")
    ap.add_argument(
        "--root-index", type=Path, default=INDEX_DIR, help="Top folder for all projects"
    )
    ap.add_argument(
        "--disable-legal", action="store_true", 
        help="Disable legal document processing and metadata extraction"
    )
    ap.add_argument(
        "--chunk-size", type=int, default=800,
        help="Chunk size for text splitting (default: 800)"
    )
    args = ap.parse_args()
    
    ingest(
        pdf_dir=args.pdf_dir, 
        project=args.project, 
        root_index=args.root_index,
        chunk_size=args.chunk_size,
        enable_legal_processing=not args.disable_legal
    )
